

import fitz
from docx import shared
from tqdm import tqdm
import docx

rotate = int(0)
zoom_x = 1.0
zoom_y = 1.0
trans = fitz.Matrix(zoom_x, zoom_y).preRotate(rotate)
# 原pdf文件全名
open_file_path = r'C:\Users\15839\Desktop\原创力 C#网络应用编程 第3版_13619182_compressed.pdf'
# 图片存储目录
save_file_path = r'C:\Users\15839\Desktop\新建文件夹'
pdf = fitz.open(open_file_path)
# 在内存中创建一个新docx文档对象
doc = docx.Document()

docx_num = 4
j = 1
pont = []

while j < docx_num:
    pont.append(len(range(pdf.pageCount))/docx_num*j)
    j = j+1

j = 0

for i in tqdm(range(pdf.pageCount)):

    if i < pdf.pageCount / 2:
        pm = pdf[i].getPixmap(matrix=trans, alpha=False)
        png_path = save_file_path + '/%s.png' % i
        # 保存提取出的图片
        pm.writePNG(png_path)
        # 将图片插入docx文档的尾部,并设置宽度
        doc.add_picture(png_path, width=shared.Cm(15))

    elif i == pdf.pageCount / 2:
        pm = pdf[i].getPixmap(matrix=trans, alpha=False)
        png_path = save_file_path + '/%s.png' % i
        pm.writePNG(png_path)
        doc.add_picture(png_path, width=shared.Cm(15))
        # 保存docx文件,参数为文件全名
        doc.save(r'C:\Users\15839\Desktop\test1.docx')
        doc = docx.Document()

    else:
        pm = pdf[i].getPixmap(matrix=trans, alpha=False)
        png_path = save_file_path + '/%s.png' % i
        pm.writePNG(png_path)
        doc.add_picture(png_path, width=shared.Cm(15))

# 保存docx文件,参数为文件全名
doc.save(r'C:\Users\15839\Desktop\test2.docx')
