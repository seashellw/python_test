import docx
from tqdm import tqdm

start_i = 1
end_i = 61

doc_end = docx.Document()

for path_i in tqdm(range(start_i, end_i + 1)):
    path_0 = 'C:\\Users\\15839\\Desktop\\淫辱优等生-真红乐章-word\\1 (' + str(
        path_i) + ').docx'
    doc_0 = docx.Document(path_0)
    text = ''
    for item in doc_0.paragraphs:
        text = text + item.text
    doc_end.add_paragraph(text)


doc_end.save('C:\\Users\\15839\\Desktop\\淫辱优等生.docx')

print('完成')
